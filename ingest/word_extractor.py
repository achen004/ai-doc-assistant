"""Word document extraction module for .doc and .docx files."""

import os
import logging
from typing import List, Dict, Any
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
import zipfile
import tempfile
from PIL import Image
import io

# For .doc files (legacy format)
try:
    import win32com.client
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordExtractor:
    """Extract text and images from Word documents (.doc and .docx)."""
    
    def __init__(self):
        self.supported_formats = ['.doc', '.docx']
    
    def extract_text_and_images(self, file_path: str, output_dir: str = "data/images") -> Dict[str, Any]:
        """
        Extract text and images from Word document.
        
        Args:
            file_path: Path to the Word document
            output_dir: Directory to save extracted images
            
        Returns:
            Dictionary containing extracted text and images
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self._extract_from_docx(file_path, output_dir)
        elif file_ext == '.doc':
            return self._extract_from_doc(file_path, output_dir)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_docx(self, file_path: str, output_dir: str) -> Dict[str, Any]:
        """Extract from .docx files using python-docx."""
        try:
            doc = Document(file_path)
            
            # Extract text
            text_pages = []
            full_text = ""
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
            
            # Add text as single "page"
            text_pages.append({
                "file": file_path,
                "page": 1,
                "text": full_text
            })
            
            # Extract images
            image_metadata = self._extract_images_from_docx(file_path, output_dir)
            
            # Get document metadata
            props = doc.core_properties
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
                'pages': 1,  # Word docs don't have pages like PDFs
                'file_path': file_path,
                'word_count': len(full_text.split()),
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()])
            }
            
            return {
                'text_pages': text_pages,
                'image_metadata': image_metadata,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX {file_path}: {str(e)}")
            raise
    
    def _extract_images_from_docx(self, file_path: str, output_dir: str) -> List[Dict[str, Any]]:
        """Extract images from .docx file."""
        os.makedirs(output_dir, exist_ok=True)
        image_metadata = []
        
        try:
            # Open docx as zip file to access images
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Find image files in the zip
                image_files = [f for f in zip_file.namelist() 
                             if f.startswith('word/media/') and 
                             any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])]
                
                for i, img_file in enumerate(image_files):
                    try:
                        # Extract image data
                        img_data = zip_file.read(img_file)
                        
                        # Get file extension
                        original_ext = os.path.splitext(img_file)[1]
                        if not original_ext:
                            original_ext = '.png'
                        
                        # Save image
                        img_filename = f"{output_dir}/word_{os.path.basename(file_path)}_{i}{original_ext}"
                        
                        with open(img_filename, 'wb') as img_out:
                            img_out.write(img_data)
                        
                        # Get image dimensions
                        try:
                            with Image.open(io.BytesIO(img_data)) as img:
                                width, height = img.size
                        except Exception:
                            width, height = 0, 0
                        
                        image_metadata.append({
                            "file": file_path,
                            "page": 1,
                            "image_path": img_filename,
                            "image_index": i,
                            "width": width,
                            "height": height
                        })
                        
                    except Exception as e:
                        logger.error(f"Error extracting image {img_file}: {str(e)}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error accessing images in DOCX: {str(e)}")
        
        return image_metadata
    
    def _extract_from_doc(self, file_path: str, output_dir: str) -> Dict[str, Any]:
        """Extract from .doc files using win32com (Windows only)."""
        if not WIN32_AVAILABLE:
            # Fallback: try to convert to docx first
            return self._convert_doc_to_docx_and_extract(file_path, output_dir)
        
        try:
            # Use Word COM object to open .doc file
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            
            # Extract text
            full_text = doc.Content.Text
            
            text_pages = [{
                "file": file_path,
                "page": 1,
                "text": full_text
            }]
            
            # Extract images (basic implementation)
            image_metadata = []
            try:
                for i, shape in enumerate(doc.InlineShapes):
                    if shape.Type == 3:  # Picture type
                        # Save image
                        img_filename = f"{output_dir}/word_doc_{i}.png"
                        os.makedirs(output_dir, exist_ok=True)
                        
                        # This is a simplified extraction - may not work for all image types
                        try:
                            shape.Range.CopyAsPicture()
                            # More complex image extraction would be needed here
                            image_metadata.append({
                                "file": file_path,
                                "page": 1,
                                "image_path": img_filename,
                                "image_index": i,
                                "width": 0,
                                "height": 0
                            })
                        except Exception as e:
                            logger.error(f"Could not extract image {i}: {str(e)}")
            except Exception as e:
                logger.error(f"Error extracting images from DOC: {str(e)}")
            
            # Get metadata
            metadata = {
                'title': doc.BuiltInDocumentProperties("Title").Value or '',
                'author': doc.BuiltInDocumentProperties("Author").Value or '',
                'pages': 1,
                'file_path': file_path,
                'word_count': len(full_text.split())
            }
            
            doc.Close()
            word.Quit()
            
            return {
                'text_pages': text_pages,
                'image_metadata': image_metadata,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error extracting from DOC {file_path}: {str(e)}")
            if 'word' in locals():
                try:
                    word.Quit()
                except:
                    pass
            raise
    
    def _convert_doc_to_docx_and_extract(self, file_path: str, output_dir: str) -> Dict[str, Any]:
        """Fallback: Convert .doc to .docx and then extract."""
        logger.warning(f"WIN32COM not available. Attempting alternative extraction for {file_path}")
        
        # This is a simplified fallback - in production you might want to use:
        # - LibreOffice headless conversion
        # - Online conversion services
        # - Other conversion tools
        
        try:
            # Try to read as text (very basic)
            with open(file_path, 'rb') as f:
                content = f.read()
                # Very basic text extraction - this won't work well for complex docs
                text = content.decode('utf-8', errors='ignore')
                # Clean up the text (remove control characters etc.)
                text = ''.join(char for char in text if char.isprintable() or char.isspace())
                
            text_pages = [{
                "file": file_path,
                "page": 1,
                "text": text
            }]
            
            metadata = {
                'title': os.path.basename(file_path),
                'author': '',
                'pages': 1,
                'file_path': file_path,
                'word_count': len(text.split()),
                'extraction_method': 'fallback_binary_read'
            }
            
            return {
                'text_pages': text_pages,
                'image_metadata': [],  # No images in fallback mode
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Fallback extraction failed for {file_path}: {str(e)}")
            raise